from __future__ import annotations

import io
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from pathlib import Path
from typing import List

import fitz  # PyMuPDF

# PyMuPDF (MuPDF C library) keeps process-level global state — font caches,
# colour profiles, glyph atlases — that is NOT thread-safe on Windows.
# A threading.Lock only prevents *concurrent* entry; it does not protect
# against corruption that was already written to global state by a prior
# concurrent call that snuck in before the lock was added.
#
# The only reliable fix: route ALL fitz work through a single dedicated OS
# thread. ThreadPoolExecutor(max_workers=1) guarantees that, regardless of
# how many pipeline threads call these functions simultaneously, the actual
# fitz code executes serially in one thread.
_fitz_thread = ThreadPoolExecutor(max_workers=1, thread_name_prefix="fitz")


@dataclass
class PDFAnalysisResult:
    doc_type: str  # DIGITAL or SCANNED
    confidence: float
    total_pages: int
    pages_with_text: int
    text_coverage_percent: float
    total_text_length: int
    page_texts: List[str]
    full_text: str
    reason: str


def _analyze_pdf_impl(file_path: str | Path | bytes) -> PDFAnalysisResult:
    if isinstance(file_path, bytes):
        doc = fitz.open(stream=file_path, filetype="pdf")
    else:
        doc = fitz.open(str(file_path))

    total_pages = len(doc)
    page_texts = []
    pages_with_text = 0

    for page in doc:
        text = page.get_text().strip()
        page_texts.append(text)
        if len(text) > 10:
            pages_with_text += 1

    doc.close()

    full_text = "\n".join(page_texts)
    total_text_length = len(full_text.strip())
    text_coverage_percent = (pages_with_text / total_pages * 100) if total_pages > 0 else 0

    is_digital = total_text_length > 200 or (total_text_length > 100 and text_coverage_percent >= 80)

    if is_digital:
        return PDFAnalysisResult(
            doc_type="DIGITAL",
            confidence=0.95,
            total_pages=total_pages,
            pages_with_text=pages_with_text,
            text_coverage_percent=text_coverage_percent,
            total_text_length=total_text_length,
            page_texts=page_texts,
            full_text=full_text,
            reason=f"PDF has {total_text_length} chars across {pages_with_text}/{total_pages} pages",
        )
    else:
        return PDFAnalysisResult(
            doc_type="SCANNED",
            confidence=0.90,
            total_pages=total_pages,
            pages_with_text=pages_with_text,
            text_coverage_percent=text_coverage_percent,
            total_text_length=total_text_length,
            page_texts=page_texts,
            full_text=full_text,
            reason=f"Insufficient text ({total_text_length} chars, {text_coverage_percent:.0f}% coverage)",
        )


def _pdf_page_to_image_impl(file_path: str | Path | bytes, page_num: int, dpi: int) -> bytes:
    if isinstance(file_path, bytes):
        doc = fitz.open(stream=file_path, filetype="pdf")
    else:
        doc = fitz.open(str(file_path))

    page = doc[page_num]
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("png")
    doc.close()
    return img_bytes


def analyze_pdf(file_path: str | Path | bytes) -> PDFAnalysisResult:
    """
    Analyze a PDF to determine if it is digital (text-based) or scanned.
    Runs inside the dedicated single-thread fitz executor — never concurrent.
    """
    return _fitz_thread.submit(_analyze_pdf_impl, file_path).result()


def pdf_page_to_image(file_path: str | Path | bytes, page_num: int = 0, dpi: int = 300) -> bytes:
    """Convert a PDF page to PNG bytes. Runs in the single-thread fitz executor."""
    return _fitz_thread.submit(_pdf_page_to_image_impl, file_path, page_num, dpi).result()


def extract_docx_text(file_path: str | Path | bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        if isinstance(file_path, bytes):
            doc = docx.Document(io.BytesIO(file_path))
        else:
            doc = docx.Document(str(file_path))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
